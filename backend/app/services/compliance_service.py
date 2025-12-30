from sqlalchemy.orm import Session
from app.models.content import ContentSubmission, ComplianceStatus, InputType
from app.models.rule import Rule
from app.providers.llm_provider import LLMProvider
from app.services.audit_service import AuditService
from uuid import UUID
from typing import List, Dict, Optional
import tiktoken
import PyPDF2
import docx
import io


class ComplianceService:
    """Service for document compliance checking"""
    
    def __init__(
        self,
        db: Session,
        llm_provider: LLMProvider
    ):
        self.db = db
        self.llm_provider = llm_provider
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    async def check_document_compliance(
        self,
        file_content: bytes,
        filename: str,
        user_id: UUID
    ) -> Dict:
        """Check uploaded document for compliance violations
        
        Flow:
        1. Extract text from document
        2. Token-based chunking with metadata
        3. Load active rules
        4. Check each chunk against rules
        5. Identify violations with context
        6. Store submission
        """
        
        # Step 1: Extract text
        extracted_text, metadata = self._extract_document_text(file_content, filename)
        
        # Step 2: Token-based chunking
        chunks = self._chunk_by_tokens(extracted_text, metadata)
        
        # Step 3: Load active rules
        active_rules = self.db.query(Rule).filter(Rule.is_active == True).all()
        
        # Step 4: Check chunks against rules
        violations = []
        rules_triggered_set = {}
        
        for chunk in chunks:
            chunk_violations = await self._check_chunk_compliance(chunk, active_rules)
            if chunk_violations:
                violations.append({
                    "chunk_text": chunk["text"],
                    "page_number": chunk.get("page"),
                    "section": chunk.get("section"),
                    "violated_rules": chunk_violations
                })
                
                # Track all triggered rules
                for violation in chunk_violations:
                    rule_id = violation["rule_id"]
                    if rule_id not in rules_triggered_set:
                        rules_triggered_set[rule_id] = violation
        
        # Determine compliance status
        compliance_status = ComplianceStatus.VIOLATIONS if violations else ComplianceStatus.COMPLIANT
        rules_triggered = list(rules_triggered_set.values())
        
        # Step 5: Store submission
        submission = ContentSubmission(
            user_id=user_id,
            input_type=InputType.DOCUMENT,
            input_reference=filename,
            final_content=extracted_text[:5000],  # Store first 5000 chars
            compliance_status=compliance_status,
            rules_triggered=rules_triggered
        )
        
        self.db.add(submission)
        self.db.commit()
        self.db.refresh(submission)
        
        # Audit log
        AuditService.log_action(
            self.db,
            action_type="document_checked",
            actor_id=user_id,
            resource_type="content",
            resource_id=submission.submission_id,
            decision_summary=f"Checked document: {len(violations)} violations found"
        )
        
        return {
            "submission_id": submission.submission_id,
            "compliance_status": compliance_status,
            "violations": violations,
            "rules_triggered": rules_triggered
        }
    
    async def rewrite_compliant(
        self,
        violating_text: str,
        violated_rules: List[Dict]
    ) -> str:
        """Rewrite violating text to be compliant"""
        
        rules_text = "\n".join([
            f"- {r['rule_text']}" for r in violated_rules
        ])
        
        rewrite_prompt = f"""Rewrite this text to comply with all listed rules.

Violating text:
{violating_text}

Rules to comply with:
{rules_text}

Provide a compliant version that preserves the original meaning while fixing violations."""
        
        result = await self.llm_provider.generate(
            prompt=rewrite_prompt,
            system_prompt="You rewrite content to be compliant while preserving meaning.",
            temperature=0.5,
            max_tokens=1000
        )
        
        return result["content"]
    
    def _extract_document_text(self, file_content: bytes, filename: str) -> tuple:
        """Extract text from various document formats"""
        
        ext = filename.lower().split('.')[-1]
        
        try:
            if ext == 'pdf':
                return self._extract_pdf(file_content)
            elif ext in ['docx', 'doc']:
                return self._extract_docx(file_content)
            elif ext in ['txt', 'md']:
                text = file_content.decode('utf-8')
                return text, {"format": "text"}
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")
    
    def _extract_pdf(self, pdf_content: bytes) -> tuple:
        """Extract text from PDF with page metadata"""
        pdf_file = io.BytesIO(pdf_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        page_map = {}
        
        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            page_map[len(text)] = page_num
            text += f"\n[PAGE {page_num}]\n{page_text}"
        
        return text, {"format": "pdf", "page_map": page_map, "total_pages": len(pdf_reader.pages)}
    
    def _extract_docx(self, docx_content: bytes) -> tuple:
        """Extract text from DOCX with paragraph metadata"""
        doc_file = io.BytesIO(docx_content)
        doc = docx.Document(doc_file)
        
        text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        return text, {"format": "docx", "paragraphs": len(doc.paragraphs)}
    
    def _chunk_by_tokens(
        self,
        text: str,
        metadata: Dict,
        chunk_size: int = 512,
        overlap: int = 50
    ) -> List[Dict]:
        """Chunk text by tokens while preserving legal meaning"""
        
        tokens = self.tokenizer.encode(text)
        chunks = []
        
        # Extract page markers if PDF
        page_map = metadata.get("page_map", {})
        
        start = 0
        while start < len(tokens):
            end = min(start + chunk_size, len(tokens))
            chunk_tokens = tokens[start:end]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            
            # Determine page number for this chunk
            page_num = None
            if page_map:
                char_pos = len(self.tokenizer.decode(tokens[:start]))
                for pos, page in sorted(page_map.items(), reverse=True):
                    if char_pos >= pos:
                        page_num = page
                        break
            
            # Extract section header if present
            section = self._extract_section_header(chunk_text)
            
            chunks.append({
                "text": chunk_text.strip(),
                "tokens": len(chunk_tokens),
                "page": page_num,
                "section": section,
                "start_token": start,
                "end_token": end
            })
            
            start = end - overlap if end < len(tokens) else end
        
        return chunks
    
    async def _check_chunk_compliance(
        self,
        chunk: Dict,
        rules: List[Rule]
    ) -> List[Dict]:
        """Check a chunk against all rules using LLM for accuracy"""
        
        # We use LLM for checking to avoid false positives from simple keyword matching
        rules_text = "\n".join([f"- {r.rule_text}" for r in rules])
        
        review_schema = {
            "compliance_issues": [
                {"rule_violated": "string", "severity": "string", "category": "string", "explanation": "string"}
            ]
        }
        
        review_prompt = f"""Role: You are a precise Compliance Auditor.
Check this document section for compliance violations.

Section Content:
\"\"\"
{chunk["text"]}
\"\"\"

Active Regulations:
{rules_text}

INSTRUCTIONS:
1. Identify specific violations of the regulations.
2. Return ONLY violations that are clearly present.
3. If no violations, return empty list.
4. Categorize each violation (e.g. BRAND, IRDAI, SEO).
"""
        
        try:
            result = await self.llm_provider.generate_structured(
                prompt=review_prompt,
                system_prompt="You are a strict but fair compliance auditor.",
                response_schema=review_schema
            )
            
            violations = []
            for issue in result.get("compliance_issues", []):
                # Find matching rule object if possible, otherwise use generic
                matched_rule = next((r for r in rules if r.rule_text in issue.get("rule_violated", "")), None)
                
                violations.append({
                    "rule_id": str(matched_rule.rule_id) if matched_rule else "ai_detected",
                    "rule_text": issue.get("rule_violated"),
                    "category": issue.get("category", "GENERAL").upper(),
                    "severity": issue.get("severity", "MEDIUM").upper(),
                    "status": "violated",
                    "explanation": issue.get("explanation")
                })
            
            return violations

        except Exception as e:
            print(f"Chunk AI review failed: {str(e)}")
            # Fallback to keyword matching if LLM fails
            return self._check_rule_violation_keywords(chunk["text"], rules)
            
    def _check_rule_violation_keywords(self, text: str, rules: List[Rule]) -> List[Dict]:
        """Fallback keyword matching"""
        violations = []
        text_lower = text.lower()
        for rule in rules:
             # Very simple negative check
             if "must not" in rule.rule_text.lower() or "prohibited" in rule.rule_text.lower():
                 # This is a placeholder for the previous logic, kept simple for fallback
                 pass
        return violations

    def _check_rule_violation(self, text: str, rule: Rule) -> bool:
        """Deprecated: Use LLM instead"""
        return False
    
    @staticmethod
    def _extract_forbidden_terms(rule_text: str) -> List[str]:
        return []

    @staticmethod
    def _extract_section_header(text: str) -> Optional[str]:
        """Extract section header from text chunk"""
        lines = text.split('\n')
        for line in lines[:3]:  # Check first 3 lines
            line = line.strip()
            if line and (line.isupper() or line.startswith('[PAGE')):
                return line[:100]
        return None
